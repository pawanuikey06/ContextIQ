"""
Generate VRIZE Hackathon Submission as a .docx file
matching the submission_template.docx format.

Updated: February 2026 — reflects ALL current ContextIQ features.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VRIZE Video Analytics Hackathon: Team Submission")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    "Instructions for Teams: Please complete this template and upload it to your "
    "designated OneDrive folder along with all referenced output files. Ensure every "
    "stage of your pipeline is documented clearly so the evaluation committee can "
    "assess your work easily."
).italic = True

# ── Team Information ──
h = doc.add_heading("Team Information", level=1)
doc.add_paragraph().add_run("Team Name: ").bold = True
doc.paragraphs[-1].add_run("Squad404")

doc.add_paragraph().add_run("Team Members: ").bold = True
doc.paragraphs[-1].add_run("Pawan Kumar Uikey, Ashish Jaiswal, Richa Pandey")

doc.add_paragraph().add_run("Project Name: ").bold = True
doc.paragraphs[-1].add_run(
    "ContextIQ — a fully-featured Meeting Intelligence Platform that takes raw MS Teams "
    "video recordings and produces speaker-diarized transcriptions, voice-identified speakers, "
    "bilingual AI summaries (English + Hindi), sentiment analysis, topic segmentation, "
    "action item extraction with Jira integration, SOW drafts, per-speaker report cards, "
    "and a RAG-powered chatbot — all orchestrated through a modern Svelte frontend and "
    "FastAPI backend."
)

# ══════════════════════════════════════════════════════════════
# PART 1: Open-Source Tool Registry
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 1: Open-Source Tool Registry", level=1)
doc.add_paragraph(
    "List of all open-source tools, libraries, and models used during the hackathon."
)

tools = [
    ("FFmpeg", "v7.0+", "Audio extraction from MS Teams .mp4 video (16 kHz mono WAV)", "Stage 1"),
    ("noisereduce", "Latest", "Spectral gating noise reduction preprocessing", "Stage 1"),
    ("soundfile", "Latest", "Audio file I/O for preprocessing pipeline", "Stage 1"),
    ("AssemblyAI SDK", "Latest", "Primary cloud-based STT with speaker diarization", "Stage 2"),
    ("Groq Whisper API", "Latest", "Ultra-fast cloud STT (whisper-large-v3-turbo)", "Stage 2"),
    ("WhisperX", "v3.1", "Local STT with word-level forced alignment (wav2vec2)", "Stage 2"),
    ("pyannote.audio", "v3.1", "Neural speaker diarization — identifies individual speakers", "Stage 2"),
    ("SpeechBrain (ECAPA-TDNN)", "Latest", "192-dim voice embeddings for speaker identification", "Stage 2"),
    ("PyTorch", "2.x (CUDA 12.8)", "GPU-accelerated ML inference for local models", "Stage 2"),
    ("Groq SDK (Llama 3.3 70B)", "Latest", "Ultra-fast LLM for summaries, action items, sentiment, topics", "Stage 3"),
    ("LangChain", "Latest", "RAG pipeline orchestration (retrieval, chains, memory)", "Stage 3"),
    ("ChromaDB", "Latest", "Local vector database for transcript embeddings (RAG)", "Stage 3"),
    ("HuggingFace (all-MiniLM-L6-v2)", "Latest", "384-dim embedding model for semantic search", "Stage 3"),
    ("Chart.js", "Latest", "Interactive charts for sentiment, analytics, culture score", "Stage 3"),
    ("fpdf2", "Latest", "PDF report generation with Unicode Hindi support", "Stage 4"),
    ("smtplib (stdlib)", "Built-in", "Email publishing with PDF attachments via SMTP", "Stage 4"),
    ("Atlassian REST API", "v3", "Jira integration — push, sync, update action items", "Stage 4"),
    ("FastAPI", "0.100+", "Backend REST API server with 35+ endpoints", "All"),
    ("Uvicorn", "Latest", "ASGI server to run FastAPI", "All"),
    ("Svelte", "v5", "Modern frontend SPA with reactive UI", "All"),
    ("Vite", "v5", "Frontend build tool and dev server", "All"),
    ("TailwindCSS", "v3", "Utility-first CSS framework for responsive UI", "All"),
    ("Lucide Svelte", "Latest", "Icon library for the frontend", "All"),
    ("svelte-spa-router", "Latest", "Client-side hash-based routing for SPA", "All"),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(["Tool / Library Name", "Version", "Primary Purpose", "Stage Used"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

for tool in tools:
    row = table.add_row().cells
    for i, val in enumerate(tool):
        row[i].text = val

# ══════════════════════════════════════════════════════════════
# PART 2: SOP & Pipeline Stages
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 2: Standard Operating Procedure (SOP) & Pipeline Stages", level=1)
doc.add_paragraph(
    "Breakdown of the video analytics pipeline into sequential stages."
)

# --- Stage 1 ---
doc.add_heading("Stage 1: Data Pre-Processing & Audio Extraction", level=2)

doc.add_paragraph().add_run("Objective: ").bold = True
doc.paragraphs[-1].add_run(
    "Extract clean, optimized audio from raw MS Teams .mp4 video recordings "
    "and prepare it for transcription."
)

doc.add_paragraph().add_run("Tool(s) Used: ").bold = True
doc.paragraphs[-1].add_run("FFmpeg, noisereduce, soundfile")

doc.add_paragraph().add_run("Input Data: ").bold = True
doc.paragraphs[-1].add_run("Original MS Teams .mp4 video file uploaded via the web UI.")

doc.add_paragraph().add_run("Output File Link: ").bold = True
doc.paragraphs[-1].add_run("data/audio/{meeting_id}.wav — 16 kHz mono WAV file")

doc.add_paragraph().add_run("Execution Details: ").bold = True
doc.add_paragraph(
    "1. User uploads video via POST /upload-video endpoint.\n"
    "2. SHA-256 hash is computed to prevent duplicate processing.\n"
    "3. FFmpeg extracts audio: ffmpeg -i input.mp4 -ar 16000 -ac 1 -f wav output.wav\n"
    "4. Audio preprocessing applies spectral gating noise reduction (noisereduce) and peak normalization.\n"
    "5. Clean audio saved as {meeting_id}_clean.wav.",
    style='List Bullet'
)

# --- Stage 2 ---
doc.add_heading("Stage 2: Transcription, Diarization & Voice Identification", level=2)

doc.add_paragraph().add_run("Objective: ").bold = True
doc.paragraphs[-1].add_run(
    "Convert the extracted audio to text with word-level timestamps, identify individual "
    "speakers (diarization), and automatically recognize known speakers using voice embeddings."
)

doc.add_paragraph().add_run("Tool(s) Used: ").bold = True
doc.paragraphs[-1].add_run(
    "AssemblyAI (primary), Groq Whisper (ultra-fast), WhisperX + pyannote.audio (local), "
    "SpeechBrain ECAPA-TDNN (voice identification)"
)

doc.add_paragraph().add_run("Input Data: ").bold = True
doc.paragraphs[-1].add_run("data/audio/{meeting_id}_clean.wav (preprocessed audio)")

doc.add_paragraph().add_run("Output File Links: ").bold = True
doc.add_paragraph("storage/{meeting_id}/transcript.json — full diarized transcript", style='List Bullet')
doc.add_paragraph("storage/{meeting_id}/speaker_clips/ — 10-second WAV clips per speaker", style='List Bullet')
doc.add_paragraph("storage/speaker_profiles/profiles.json — 192-dim voice embeddings for known speakers", style='List Bullet')

doc.add_paragraph().add_run("Execution Details: ").bold = True
doc.add_paragraph(
    "1. Triggered via POST /transcribe/{meeting_id} with configurable STT engine.\n"
    "2. AssemblyAI mode (primary): Audio uploaded with speaker_labels=True, returns speaker-tagged segments.\n"
    "3. Groq mode (ultra-fast): Whisper large-v3-turbo transcription + local pyannote diarization.\n"
    "4. Local mode: WhisperX with wav2vec2 forced alignment + pyannote.audio 3.1 diarization on GPU.\n"
    "5. Voice Identification: Extracts ~10s speaker clips (SNR-ranked), generates 192-dim ECAPA-TDNN embeddings, "
    "matches against stored profiles using cosine similarity (threshold 0.55).\n"
    "6. Matched speakers are auto-renamed in the transcript (e.g., SPEAKER_00 → 'Babuji Abraham').\n"
    "7. GPU memory explicitly cleared after processing.",
    style='List Bullet'
)

# --- Stage 3 ---
doc.add_heading("Stage 3: Analytics & Feature Generation", level=2)

doc.add_paragraph().add_run("Objective: ").bold = True
doc.paragraphs[-1].add_run(
    "Generate comprehensive AI-powered analytics: bilingual summaries, action items, "
    "sentiment analysis, topic segmentation, requirements, documentation, speaker report cards, "
    "SOW drafts, and enable a RAG chatbot."
)

doc.add_paragraph().add_run("Tool(s) Used: ").bold = True
doc.paragraphs[-1].add_run("Groq API (Llama 3.3 70B), LangChain, ChromaDB, HuggingFace Embeddings")

doc.add_paragraph().add_run("Input Data: ").bold = True
doc.paragraphs[-1].add_run("storage/{meeting_id}/transcript.json")

doc.add_paragraph().add_run("Output File Links: ").bold = True
outputs = [
    "storage/{meeting_id}/summary.json — Bilingual summaries (English + Hindi)",
    "storage/{meeting_id}/action_items.json — Action items, decisions, takeaways, risks",
    "storage/{meeting_id}/sentiment.json — Per-segment sentiment scores and emotion labels",
    "storage/{meeting_id}/topics.json — Topic segmentation with time ranges and speakers",
    "storage/{meeting_id}/requirements.json — Functional/non-functional requirements, user stories",
    "storage/{meeting_id}/documentation.json — Auto-generated meeting minutes (MoM)",
    "storage/{meeting_id}/followup_email.json — AI-drafted follow-up email",
    "storage/chroma_db/ — Vector embeddings indexed in ChromaDB for RAG",
]
for o in outputs:
    doc.add_paragraph(o, style='List Bullet')

doc.add_paragraph().add_run("Execution Details: ").bold = True
doc.add_paragraph(
    "1. Summarization (POST /summarize/{id}): Groq Llama 3.3 70B generates speaker-wise + overall "
    "summaries in English and Hindi.\n"
    "2. Action Items (POST /meeting/{id}/action-items): Structured JSON extraction with task, assignee, "
    "deadline, priority, category, success criteria, and dependencies.\n"
    "3. Sentiment (POST /meeting/{id}/sentiment): Per-segment mood scoring with confidence scores.\n"
    "4. Topic Segmentation (POST /meeting/{id}/topics): Identifies distinct discussion topics with "
    "time ranges, titles, summaries, and participating speakers.\n"
    "5. Requirements Mining (POST /meeting/{id}/requirements): Extracts functional/non-functional "
    "requirements, technical constraints, and user stories with MoSCoW prioritization.\n"
    "6. Speaker Report Cards (GET /meeting/{id}/speaker-analytics): Per-speaker scorecards with role "
    "classification (Decision Maker, Presenter, etc.), talk-time, and sentiment trends.\n"
    "7. Meeting Culture Score (GET /stats/culture-score): Composite health metric (0-100) combining "
    "speaker balance (Gini-like), sentiment, action completion, and meeting efficiency.\n"
    "8. RAG Indexing (POST /chat/index/{id}): all-MiniLM-L6-v2 embeddings stored in ChromaDB.\n"
    "9. RAG Chat (POST /chat/ask/stream): Diverse retrieval algorithm (round-robin across meetings) + "
    "SSE streaming with source citations.\n"
    "10. HITL Regeneration: When speakers are renamed, all 8 AI insights regenerate in background "
    "with mapped names (force=True).",
    style='List Bullet'
)

# --- Stage 4 ---
doc.add_heading("Stage 4: Report Generation & Publishing", level=2)

doc.add_paragraph().add_run("Objective: ").bold = True
doc.paragraphs[-1].add_run(
    "Generate professional PDF reports, distribute via email and Teams, "
    "push action items to Jira, and export subtitles."
)

doc.add_paragraph().add_run("Tool(s) Used: ").bold = True
doc.paragraphs[-1].add_run("fpdf2, smtplib, Microsoft Teams Webhook, Atlassian Jira REST API")

doc.add_paragraph().add_run("Input Data: ").bold = True
doc.paragraphs[-1].add_run("summary.json, action_items.json, requirements.json, documentation.json")

doc.add_paragraph().add_run("Output File Links: ").bold = True
doc.add_paragraph("PDF summary report (NotoSans + NotoSansDevanagari fonts)", style='List Bullet')
doc.add_paragraph("Full comprehensive report PDF (summary + actions + requirements + docs)", style='List Bullet')
doc.add_paragraph("SRT and VTT subtitle files with speaker labels", style='List Bullet')

doc.add_paragraph().add_run("Execution Details: ").bold = True
doc.add_paragraph(
    "1. PDF Generation (GET /publish/{id}/pdf): fpdf2 with Unicode Hindi support.\n"
    "2. Full Report (GET /publish/{id}/full-report): All analytics combined into one PDF.\n"
    "3. Email (POST /publish/{id}): SMTP with PDF attachment via Gmail App Passwords.\n"
    "4. Teams (POST /publish/{id}): Adaptive Card v1.4 with summary, decisions, action items.\n"
    "5. Jira Push (POST /meeting/{id}/jira/push): Action items as Jira tickets with ADF description, "
    "priority mapping, and bi-directional sync via transitions API.\n"
    "6. Subtitles (GET /meeting/{id}/subtitles/srt and /vtt): Standard subtitle export.",
    style='List Bullet'
)

# ══════════════════════════════════════════════════════════════
# PART 3: Features & Innovation Summary
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 3: Features & Innovation Summary", level=1)
doc.add_paragraph("Highlight the capabilities of the final analytics report.")

doc.add_heading("List of Features Built:", level=2)
features = [
    "Multi-Engine Speech-to-Text (AssemblyAI, Groq Whisper, local WhisperX) — configurable per meeting",
    "Speaker Diarization with GPU acceleration (pyannote.audio 3.1, neural VAD + clustering)",
    "Voice Identification — ECAPA-TDNN 192-dim embeddings, SNR-ranked clip selection, cosine similarity matching, "
    "profile averaging for multi-session stability",
    "Bilingual AI Summaries — Speaker-wise + overall summaries in English and Hindi (Devanagari)",
    "Action Item Extraction with assignee, deadline, priority, category, dependencies, and success criteria",
    "Decisions Tracker — who decided what, why, alternatives considered, and impact",
    "Key Takeaways — bullet-point highlights from every meeting",
    "Auto Meeting Title generation from transcript content",
    "Follow-Up Email Draft — professional email combining summary + action items, ready to send via SMTP",
    "Per-Segment Sentiment Analysis with emotion labels, confidence scores, and speaker trends",
    "Topic Segmentation — auto-detected discussion chapters with time ranges, titles, and speakers",
    "Requirements Mining — functional, non-functional, technical constraints, user stories with MoSCoW priority",
    "Documentation Generation — auto-generated meeting minutes with objective, attendees, and next steps",
    "Per-Speaker Report Cards — role classification (Decision Maker, Presenter, etc.), talk-time stats, sentiment",
    "Meeting Culture Score — composite health metric (speaker balance, sentiment, completion, efficiency)",
    "RAG Chatbot with SSE Streaming — cross-meeting Q&A with diverse retrieval (round-robin), source citations, "
    "session memory, and meeting calendar context",
    "Human-in-the-Loop (HITL) — speaker mapping triggers async regeneration of all 8 AI insights, "
    "summary editing & approval before publishing",
    "Professional PDF Reports with Unicode Hindi support (NotoSans + NotoSansDevanagari)",
    "SRT/VTT Subtitle Export with speaker labels",
    "One-Click Email & Teams Publishing (Adaptive Card v1.4 with rich formatting)",
    "Bidirectional Jira Integration — push tickets, sync statuses, transitions API, ADF descriptions",
    "Confluence & Notion Integration — push meeting notes to enterprise wikis via API",
    "Upload Deduplication via SHA-256 hashing (prevents reprocessing identical files)",
    "Meeting Dashboard with unique speaker counting (resolves mapped names), real-time stats, and search",
    "Global Keyword Search across transcripts, titles, and speakers with weighted relevance scoring",
    "Video Playback — in-browser video player with HTTP Range request support for seeking",
    "Skeleton Loading & Toast Notifications for polished, premium UX",
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("Innovation Highlight:", level=2)
doc.add_paragraph(
    "ContextIQ's most innovative aspect is its end-to-end, privacy-first architecture "
    "that combines local GPU-accelerated processing for speech-to-text with cloud LLM APIs "
    "for intelligence. Unlike competitors that require all data to be uploaded to external "
    "servers, ContextIQ keeps raw audio and transcripts local while only sending text to "
    "Groq for analysis."
)
doc.add_paragraph(
    "Key innovations include: (1) Voice Identification using ECAPA-TDNN neural embeddings — "
    "speakers enrolled from one meeting are automatically recognized in all future meetings; "
    "(2) Bilingual Hindi summaries in native Devanagari (absent from all major competitors); "
    "(3) Bi-directional Jira integration using the transitions API for proper status changes; "
    "(4) A diverse retrieval algorithm for the RAG chatbot that round-robins chunks across meetings "
    "to prevent single-meeting bias; (5) Human-in-the-Loop workflow where renaming speakers triggers "
    "async regeneration of all 8 AI insights; (6) Meeting Culture Score — a composite health metric "
    "combining Gini-like speaker balance, sentiment health, action completion rate, and decisions-per-minute."
)

# ══════════════════════════════════════════════════════════════
# PART 4: Final Deliverables
# ══════════════════════════════════════════════════════════════
doc.add_heading("Part 4: Final Deliverables", level=1)

doc.add_paragraph().add_run("Link to Final Analytics Report: ").bold = True
doc.paragraphs[-1].add_run("[Insert OneDrive link to the final PDF report]")

doc.add_paragraph().add_run("Link to Source Code Repository: ").bold = True
doc.paragraphs[-1].add_run("https://github.com/pawanuikey06/ContextIQ")

doc.add_paragraph().add_run("Self-Reported Transcript Accuracy (Optional): ").bold = True
doc.paragraphs[-1].add_run("[Insert percentage match and metric used, e.g., Word Error Rate]")

# ── Save ──
output_path = r"d:\ContextIQ\VRIZE_Submission_Squad404.docx"
doc.save(output_path)
print(f"✅ Submission saved to: {output_path}")
